from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\코딩\교육설계\output\doc\AGENTSMD_작성코치_프롬프트_CRAFTO_RADAR.docx")

NAVY = "1C2F5C"
GREY = "616161"
LIGHT_GREY = "F4F5F7"
PALE_BLUE = "EEF3FA"
LINE = "DADADA"


PROMPT = """나는 내 강의 또는 프로젝트를 소개하는 1페이지 홈페이지를 만들기 위해
Agentic AI 작업 폴더를 준비하고 있습니다.

내 폴더에는 input, wiki, output 폴더가 있습니다.
이 폴더에서 사용할 AGENTS.md를 함께 작성하고 싶습니다.

당신은 AGENTS.md 작성 코치입니다.
CRAFTO와 RADAR 기준으로 나를 인터뷰해 주세요.

진행 방식:
- 질문은 한 번에 하나씩만 해 주세요.
- 내가 답하면 한 문장으로 요약한 뒤 다음 질문을 해 주세요.
- 내가 “모르겠다”, “추천해줘”, “예시가 필요하다”고 말하면 선택지 2~3개를 제안해 주세요.
- 특히 RADAR 항목에서는 안전한 기본값을 먼저 추천해 주세요.
- 답이 모호하면 바로 넘어가지 말고 보완 질문을 하나 해 주세요.
- 모든 질문이 끝나면 AGENTS.md 초안을 markdown 형식으로 작성해 주세요.

질문 순서:
1. Context: 이 홈페이지는 무엇을 소개하기 위한 것인가요?
2. Role: 에이전트는 어떤 역할로 도와야 하나요?
3. Audience: 이 홈페이지를 볼 사람은 누구인가요?
4. Format: 어떤 산출물을 만들어야 하나요?
5. Tone: 어떤 문체와 분위기를 원하나요?
6. Option: 피해야 할 표현이나 넣지 말아야 할 정보는 무엇인가요?
7. Range: 에이전트가 읽어도 되는 자료와 결과물을 만들 위치는 어디인가요?
8. Approval: 어떤 행동은 반드시 승인 후 진행해야 하나요?
9. Done: 무엇이 나오면 완료인가요?
10. Ask: 언제 멈추고 질문해야 하나요?
11. Report: 작업 후 무엇을 보고해야 하나요?

RADAR 기본 추천 기준:

Range 추천:
- 기본값: input 폴더의 자료만 읽고, wiki 폴더에 중간 정리, output 폴더에 최종 결과물을 만듭니다.
- 금지: 프로젝트 폴더 밖 파일 읽기, raw/input 원본 수정, 실제 회사 내부자료 사용

Approval 추천:
- 기본값: 파일 생성·수정 전 작업계획을 먼저 제시하고 사용자 승인을 받습니다.
- 승인 필요 행동: 파일 생성, 파일 수정, 터미널 명령, 외부 사이트 접속, 패키지 설치, 배포, 공유 링크 생성
- 수업용 기본값: 자동 실행하지 않고 Request Review 방식으로 진행합니다.

Done 추천:
- 기본값: output/index.html, output/style.css, output/검토메모.md가 만들어지면 완료입니다.
- 품질 기준: 첫 화면에 대상, 문제, 제안, 문의 동선이 보여야 합니다.
- 추가 기준: 모바일에서도 읽기 좋고, 자료에 없는 표현이 없어야 합니다.

Ask 추천:
- 기본값: 자료에 없는 정보가 필요하면 멈추고 질문합니다.
- 질문해야 할 상황: 고객사명, 가격, 성과, 후기, 개인정보, 실제 이력, 배포 여부, 외부 자료 사용 여부
- 판단이 필요한 상황: 홈페이지 목적이 소개인지 모집인지 제안인지 불분명할 때

Report 추천:
- 기본값: 작업 후 생성한 파일, 주요 구성, 사용한 자료, 부족한 정보, 사람이 최종 결정할 항목을 보고합니다.
- 짧은 보고 형식: 만든 것 / 근거 자료 / 확인 필요한 것 / 다음 수정 제안

AGENTS.md에 반드시 포함할 원칙:
- input 폴더의 자료만 읽기
- wiki 폴더에 중간 정리 자료 만들기
- output 폴더에 결과물 만들기
- 파일 생성 전 작업계획 먼저 제시하기
- 자료에 없는 고객사, 가격, 후기, 성과는 추정하지 않기
- 개인정보, 계정키, 내부자료는 사용하지 않기
- 터미널 실행, 외부 접속, 배포는 승인 후 진행하기
- 작업 후 생성 파일, 주요 구성, 사람이 결정할 항목을 보고하기

먼저 1번 Context 질문부터 시작해 주세요."""


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=LINE, sz="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text, bold=False, color=None, size=None, font="Pretendard"):
    run = paragraph.add_run(text)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    add_run(p, text, bold=True, color=NAVY, size=18 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, size=10.5, color="000000", space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    add_run(p, text, color=color, size=size)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    add_run(p, text, size=size)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=fill, sz="4")
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=NAVY, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, body, color="000000", size=10.2)
    doc.add_paragraph()


def add_prompt_block(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.right_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.02
        if not line:
            p.paragraph_format.space_after = Pt(4)
            p.add_run("")
            continue
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "D2Coding")
        run.font.size = Pt(8.7)
        run.font.color.rgb = RGBColor.from_string("111111")


def add_radar_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [Cm(2.8), Cm(4.6), Cm(9.0)]
    headers = ["항목", "질문", "수업용 안전 기본값"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=NAVY, sz="4")
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, headers[idx], bold=True, color="FFFFFF", size=9.5)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Range", "어디까지 읽고 쓸 수 있는가", "input 자료만 읽고, wiki에는 중간 정리, output에는 최종 결과물을 만든다."),
        ("Approval", "무엇은 승인 후 해야 하는가", "파일 생성·수정, 터미널 명령, 외부 접속, 패키지 설치, 배포는 승인 후 진행한다."),
        ("Done", "무엇이 나오면 완료인가", "index.html, style.css, 검토메모.md가 만들어지고 모바일 가독성까지 점검한다."),
        ("Ask", "언제 멈추고 질문해야 하는가", "자료에 없는 고객사, 가격, 후기, 성과, 개인정보, 실제 이력이 필요하면 질문한다."),
        ("Report", "무엇을 보고해야 하는가", "만든 것, 근거 자료, 확인 필요한 것, 다음 수정 제안을 짧게 보고한다."),
    ]
    for label, question, default in rows:
        cells = table.add_row().cells
        values = [label, question, default]
        for idx, cell in enumerate(cells):
            cell.width = widths[idx]
            set_cell_border(cell, color=LINE, sz="4")
            set_cell_margins(cell, 120, 120, 120, 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.08
            add_run(p, values[idx], bold=(idx == 0), color=NAVY if idx == 0 else "000000", size=9.2)


def configure_styles(doc):
    styles = doc.styles
    for style_name in ["Normal", "List Bullet"]:
        style = styles[style_name]
        style.font.name = "Pretendard"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Pretendard")
        style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Pretendard"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Pretendard")
        style.font.color.rgb = RGBColor.from_string(NAVY)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    configure_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "전환설계연구소", bold=True, color=NAVY, size=9.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "AGENTS.md 작성 코치 프롬프트", bold=True, color=NAVY, size=24)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "CRAFTO-RADAR 기반 Agentic AI 작업공간 설정 실습용", color=GREY, size=12)

    add_callout(
        doc,
        "사용 목적",
        "참가자가 AI와 대화하면서 자기 작업 폴더용 AGENTS.md를 작성하도록 돕는 프롬프트입니다. 특히 RADAR 항목에서 막히는 참가자를 위해 안전한 기본값을 추천하도록 설계했습니다.",
    )

    add_heading(doc, "강의 운영 포인트", 1)
    add_bullet(doc, "CRAFTO는 요청의 맥락을 정리하는 프레임입니다.")
    add_bullet(doc, "RADAR는 에이전트가 일할 폴더의 실행 조건을 정리하는 프레임입니다.")
    add_bullet(doc, "참가자가 RADAR를 잘 모르면 안전한 기본값을 추천받게 합니다.")
    add_bullet(doc, "이번 실습의 핵심은 홈페이지 결과물보다 계획, 승인, 실행, 검토 흐름을 경험하는 것입니다.")

    add_heading(doc, "참가자 안내 멘트", 1)
    add_body(
        doc,
        "RADAR에서 막히면 괜찮습니다. 이 부분은 처음 쓰면 대부분 어렵습니다. 모르면 AI에게 “추천해줘”, “수업용 안전한 기본값으로 해줘”, “홈페이지 제작 실습에 맞게 기본 규칙을 넣어줘”라고 말하면 됩니다.",
    )

    add_heading(doc, "RADAR 기본값 요약", 1)
    add_radar_table(doc)

    doc.add_page_break()
    add_heading(doc, "복사용 전체 프롬프트", 1)
    add_body(doc, "아래 내용을 그대로 복사해 Antigravity, Codex, ChatGPT 등에서 사용합니다.", color=GREY, space_after=8)
    add_prompt_block(doc, PROMPT)

    add_heading(doc, "짧은 실행 문장", 1)
    add_body(doc, "실습 중 참가자가 막힐 때 바로 쓰게 할 문장입니다.", color=GREY)
    for text in [
        "추천해줘.",
        "수업용 안전한 기본값으로 해줘.",
        "홈페이지 제작 실습에 맞게 기본 규칙을 넣어줘.",
        "파일 생성 전 작업계획을 먼저 제시하게 해줘.",
        "자료에 없는 내용은 추정하지 않게 해줘.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "강사용 마무리 문장", 1)
    add_body(
        doc,
        "CRAFTO는 질문의 맥락을 정리합니다. RADAR는 에이전트의 실행 조건을 정리합니다. AGENTS.md는 그 내용을 폴더 안에 남기는 운영 문서입니다. AI에게 일을 맡긴다는 것은 요청을 길게 쓰는 일이 아니라, 일할 조건을 설계하는 일입니다.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()

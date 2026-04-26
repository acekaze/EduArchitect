from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "생성형AI_기초활용역량_2026_4시간_커리큘럼표.docx"


PRIMARY = RGBColor(31, 78, 121)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(96, 108, 120)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor = TEXT, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT


def add_body(doc: Document, text: str, *, color: RGBColor = TEXT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(10.5)
    run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.bold = True
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = PRIMARY if level == 1 else TEXT


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("생성형 AI 기초활용역량\n2026 버전 4시간 커리큘럼 제안")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("LLM 자유 활용에서 Agentic AI 이해까지")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    intro_box = doc.add_table(rows=1, cols=1)
    intro_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro_box.autofit = False
    intro_box.columns[0].width = Cm(16.2)
    cell = intro_box.cell(0, 0)
    set_cell_shading(cell, "EEF4FB")
    set_cell_text(
        cell,
        "이 과정은 생성형 AI 기능을 많이 보여주는 입문 교육이 아니라, "
        "AI를 결과물 요청 도구에서 판단을 돕는 파트너로 전환해 쓰도록 돕는 실무형 과정입니다.",
        size=11,
        color=PRIMARY,
    )

    add_heading(doc, "1. 과정 방향")
    add_body(
        doc,
        "이번 학습자 집단은 반도체·소재·장비 산업 중심의 중·고경력 실무자와 의사결정자가 함께 있는 혼합형 집단이다. "
        "따라서 도구 소개보다 업무 맥락과 판단 기준을 먼저 세우고, 기초 개념은 쉽게 설명하되 얕지 않게 다루는 설계가 필요하다."
    )
    add_bullet(doc, "기초 개념은 쉽게 설명하되 얕지 않게 다룬다.")
    add_bullet(doc, "도구 소개보다 업무 맥락과 판단 기준을 먼저 세운다.")
    add_bullet(doc, "Agentic AI는 무리한 실습보다 구조 이해와 도입 판단까지 연결한다.")

    add_heading(doc, "2. 과정 목표")
    goals = [
        "생성형 AI, 딥리서치, 소스 기반 AI, Agentic AI의 차이를 구조적으로 이해한다.",
        "LLM을 단순 질의응답이 아니라 문서 정리, 비교, 검토, 초안 작성, 의사결정 보조 수준까지 활용하는 기준을 익힌다.",
        "AI 활용에서 사람이 맡아야 할 역할과 최종 판단의 책임 범위를 명확히 구분한다.",
        "Agentic AI의 최신 흐름을 이해하고, 우리 업무에서 어디부터 도입 가능한지 판단할 수 있게 한다.",
    ]
    for goal in goals:
        add_bullet(doc, goal)

    add_heading(doc, "3. 핵심 메시지")
    messages = [
        "AI를 잘 쓰는 사람은 프롬프트를 길게 쓰는 사람이 아니라, 무엇을 물어야 하는지 아는 사람이다.",
        "AI의 진짜 가치는 답을 대신 내는 데 있지 않고, 판단에 필요한 구조를 더 빨리 드러내는 데 있다.",
        "업무 성과는 질문, 맥락, 검증, 최종 선택이 연결될 때 나온다.",
        "Agentic AI는 마법이 아니라, 이미 가능한 LLM 활용 위에 계획, 도구 사용, 검증 흐름이 얹힌 다음 단계다.",
    ]
    for message in messages:
        add_bullet(doc, message)

    add_heading(doc, "4. 권장 3단계 학습 구조")
    phases = [
        "1단계. 기초 개념과 사용 관점",
        "2단계. LLM 자유 활용과 소스 기반 실무 적용",
        "3단계. Agentic AI 이해와 도입 판단",
    ]
    for phase in phases:
        add_bullet(doc, phase)

    add_heading(doc, "5. 4시간 상세 커리큘럼")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.7), Cm(1.9), Cm(3.0), Cm(5.0), Cm(4.2), Cm(3.4)]
    headers = ["시간", "단계", "대주제", "핵심 내용", "실습 및 활동", "기대 효과"]
    for idx, (header, width) in enumerate(zip(headers, widths)):
        table.columns[idx].width = width
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9E8F5")
        set_cell_text(cell, header, bold=True, color=PRIMARY, size=10)

    rows = [
        [
            "0:00-0:30",
            "1단계",
            "왜 AI를 써도 결정은 더 어려워지는가",
            "생성형 AI 활용의 본질 이해, 결과 요청과 판단 보조의 차이, AI 시대에 더 중요해지는 인간의 역할",
            "내 업무에서 AI를 쓰며 답답했던 순간 적기, 빨라졌지만 확신은 줄어든 일 점검",
            "AI를 기능이 아니라 판단 구조의 도구로 보게 됨",
        ],
        [
            "0:30-1:10",
            "1단계",
            "생성형 AI의 기초 개념과 한계",
            "LLM, 검색형 AI, 딥리서치, 소스 기반 AI, Agentic AI의 차이, 환각과 과신, 사람이 맡아야 할 시작과 끝",
            "같은 질문을 검색, 일반 LLM, 소스 기반 도구에 각각 던졌을 때의 차이 비교",
            "AI 도구를 목적별로 구분하게 됨",
        ],
        [
            "1:10-2:00",
            "2단계",
            "LLM 자유 활용의 핵심 구조",
            "CRAFT-O 기반 질문 구조화, 질문의 역추산, 좋은 답보다 좋은 판단 재료를 얻는 질문법, 10-80-10 법칙",
            "보고, 분석, 회의 정리, 설명 자료, 고객 대응 중 자기 업무 질문 1개 재설계",
            "단발 질문에서 벗어나 재사용 가능한 질문 구조를 익힘",
        ],
        [
            "2:00-2:40",
            "2단계",
            "소스 기반 활용과 딥리서치",
            "파일 업로드 기반 활용, NotebookLM류 도구의 강점, 기술 트렌드와 문서를 찾기보다 검토하고 비교하는 방법",
            "기술 이슈 또는 업무 자료를 넣고 핵심 쟁점, 비교 포인트, 리스크 질문 뽑기",
            "자료 기반으로 더 안전하게 AI를 쓰는 감각을 익힘",
        ],
        [
            "2:40-3:20",
            "2단계",
            "선택을 끝내는 실무 프레임",
            "1-3-1 법칙, 4-View 메타 피드백, 결과물 다듬기보다 의사결정 정교화",
            "AI에게 대안 3개를 내게 하고 최적안 1개를 선택하게 한 뒤 4개 관점으로 재검토",
            "AI 답변을 소비하지 않고 검토하고 압축하는 힘이 생김",
        ],
        [
            "3:20-4:00",
            "3단계",
            "Agentic AI의 현재와 다음 단계",
            "계획, 실행, 도구 사용, 검증 흐름 이해, 일반 LLM 활용과의 차이, 업무 적용 예시, 도입 시 보안과 승인 체계",
            "여러 단계를 수행하는 에이전트형 활용 시연 또는 사례 walkthrough, 자기 업무 적용 시나리오 정리",
            "Agentic AI를 과장 없이 이해하고 도입 가능 지점을 판단함",
        ],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)

    doc.add_paragraph()
    add_heading(doc, "6. 모듈별 운영 포인트")
    points = [
        ("오프닝", "출발 질문은 AI로 무엇을 만들 수 있는가가 아니라, 왜 AI를 써도 결정은 더 어려워졌는가로 잡는다."),
        ("기초 개념", "용어를 나열하기보다 일반 LLM, 딥리서치, 소스 기반 AI, Agentic AI의 차이를 구조 중심으로 설명한다."),
        ("LLM 자유 활용", "잘 물어보는 법보다 업무를 더 잘 끝내는 법에 초점을 맞춘다."),
        ("소스 기반 활용", "기술 문서와 시장 자료를 넣고 요약만 시키지 말고 쟁점, 리스크, 추가 확인 항목까지 끌어낸다."),
        ("선택과 검증", "초안 생성보다 선택을 종결하고 책임지는 판단 구조를 훈련한다."),
        ("Agentic AI", "바로 전원 실습보다 구조 이해, 시연, 적용 판단 중심으로 가져간다."),
    ]
    for title_text, body_text in points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title_text}: ")
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT
        run = p.add_run(body_text)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT

    add_heading(doc, "7. 직무별 실습 예시")
    examples = [
        ("영업/마케팅", "고객 미팅 전 기술자료와 시장 이슈를 바탕으로 질문 리스트와 대응 논리를 정리하라."),
        ("R&D/공정/설계", "특정 기술 이슈에 대해 원인 가설 3개와 추가 확인이 필요한 데이터 항목을 정리하라."),
        ("품질/TEST/EHS", "반복 이슈 보고를 바탕으로 공통 원인, 리스크, 우선 대응순위를 구조화하라."),
        ("전략기획/IT/인사", "최신 AI 트렌드를 우리 조직 관점에서 기회, 리스크, 도입 과제로 정리하라."),
    ]
    ex_table = doc.add_table(rows=1, cols=2)
    ex_table.style = "Table Grid"
    ex_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ex_table.autofit = False
    ex_table.columns[0].width = Cm(4.0)
    ex_table.columns[1].width = Cm(12.8)
    for idx, header in enumerate(["직무군", "추천 실습 질문"]):
        cell = ex_table.rows[0].cells[idx]
        set_cell_shading(cell, "D9E8F5")
        set_cell_text(cell, header, bold=True, color=PRIMARY, size=10)
    for role, prompt in examples:
        cells = ex_table.add_row().cells
        set_cell_text(cells[0], role, size=9)
        set_cell_text(cells[1], prompt, size=9)

    add_heading(doc, "8. 운영 제안")
    ops = [
        "강의 40, 실습 60 정도의 비중이 적합하다.",
        "도구가 답을 주는 순간보다 사람이 기준을 세우는 순간을 더 강조한다.",
        "Agentic AI는 소개를 과장하지 않는다.",
        "마무리는 오늘 배운 도구보다 내 업무에서 당장 바꿀 질문 1개를 남기는 방식이 좋다.",
    ]
    for op in ops:
        add_bullet(doc, op)

    add_heading(doc, "9. 한 줄 과정 정의")
    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(0)
    closing.paragraph_format.line_spacing = 1.35
    run = closing.add_run(
        "이 과정은 생성형 AI를 더 많이 아는 교육이 아니라, "
        "AI를 활용해 더 잘 묻고, 더 잘 검토하고, 더 나은 결정을 내리는 실무자를 만드는 과정이다."
    )
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()

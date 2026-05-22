from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "06_팀탐험_보상카드_운영매뉴얼_v1.md"
OUT_PATH = BASE / "팀탐험_보상카드_운영매뉴얼_v1.docx"
IMAGE_PATH = BASE / "팀탐험_카드이미지" / "AI생성_단품카드_v1" / "AI_팀탐험_단품카드_모음_v1.png"

FONT = "Malgun Gothic"
GREEN = RGBColor(31, 51, 45)
GOLD = RGBColor(181, 132, 55)
INK = RGBColor(37, 42, 39)
MUTED = RGBColor(92, 100, 94)
PAPER = "F7F3E8"
LIGHT_GREEN = "E8EFE9"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 9.2, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Cm(0.1)
            cell.margin_bottom = Cm(0.1)
            cell.margin_left = Cm(0.12)
            cell.margin_right = Cm(0.12)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 9.2, row_idx == 0, RGBColor(255, 255, 255) if row_idx == 0 else INK)
            if row_idx == 0:
                shade_cell(cell, "1F332D")
            elif row_idx % 2 == 1:
                shade_cell(cell, "FAF8F0")


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TEAM EXPLORATION CARD")
    set_run_font(run, 12, True, GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("팀 탐험 보상 카드\n운영 매뉴얼")
    set_run_font(run, 28, True, GREEN)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("강사용 안내 문서 v1")
    set_run_font(run, 12, False, MUTED)
    p.paragraph_format.space_after = Pt(14)

    if IMAGE_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(IMAGE_PATH), width=Cm(15.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("점수판 없이 참여 행동을 모으고, 마지막 정산으로 학습 경험을 마무리하는 강의용 보상 도구")
    set_run_font(run, 10.5, False, MUTED)
    doc.add_page_break()


def apply_styles(doc):
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10)
        style.font.color.rgb = INK

    for style_name, size, color in [
        ("Heading 1", 17, GREEN),
        ("Heading 2", 13, GREEN),
        ("Heading 3", 11, GOLD),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_quote(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GREEN)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, 10, False, GREEN)
    doc.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        raw = line.strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def add_markdown_table(doc, table_lines):
    rows = parse_table(table_lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0), color=RGBColor(255, 255, 255) if r_idx == 0 else INK)
    style_table(table)
    doc.add_paragraph()


def add_paragraph_text(doc, text_line):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text_line)
    set_run_font(run, 10, False, INK)


def build_doc():
    doc = Document()
    apply_styles(doc)
    add_cover(doc)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 1  # skip top title; cover already handles it
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=2)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].startswith("> "):
                quote_lines.append(lines[i][2:].strip())
                i += 1
            add_quote(doc, quote_lines)
            continue
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:])
            set_run_font(run, 10, False, INK)
            p.paragraph_format.space_after = Pt(2)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(re.sub(r"^\d+\. ", "", line))
            set_run_font(run, 10, False, INK)
            p.paragraph_format.space_after = Pt(2)
        else:
            add_paragraph_text(doc, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("팀 탐험 보상 카드 운영 매뉴얼 v1")
    set_run_font(run, 8, False, MUTED)

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "SK하이닉스_협력사_주니어_소프트랜딩_AI랩업_제안서.docx"

PRIMARY = RGBColor(24, 66, 117)
PRIMARY_LIGHT = "EAF1F8"
HEADER_FILL = "D9E8F5"
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(90, 98, 108)


def set_korean_font(run, size: float, *, bold: bool = False, color: RGBColor = TEXT) -> None:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor = TEXT, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_korean_font(run, size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_korean_font(run, 15 if level == 1 else 12, bold=True, color=PRIMARY if level == 1 else TEXT)


def add_body(doc: Document, text: str, *, color: RGBColor = TEXT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_korean_font(run, 10.5, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_korean_font(run, 10.3)


def add_label_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(f"{label}: ")
    set_korean_font(run, 10.3, bold=True)
    run = p.add_run(body)
    set_korean_font(run, 10.3)


def add_intro_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PRIMARY_LIGHT)
    set_cell_text(cell, text, bold=False, color=PRIMARY, size=10.8)


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("SK하이닉스 협력사 주니어/영입 구성원\n소프트랜딩 과정 AI 랩업 제안서")
    set_korean_font(run, 19, bold=True, color=PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("반도체 학습 내용을 협업 가능한 결과물로 전환하는 3일 과정 설계안")
    set_korean_font(run, 10.8, color=MUTED)

    add_intro_box(
        doc,
        "이번 제안은 AI를 별도의 툴 체험으로 분리하지 않고, 앞서 학습한 반도체 시장 트렌드와 공정 이해를 "
        "NotebookLM 기반 학습 자산으로 전환해 실용성과 협업 메시지를 함께 남기는 방향으로 설계한 안입니다.",
    )

    add_heading(doc, "1. 제안 배경")
    add_body(
        doc,
        "이번 과정은 협력사 주니어/영입 구성원을 대상으로 진행되는 소프트랜딩 프로그램으로, "
        "참여자의 직무와 배경이 다양할 가능성이 높습니다. 따라서 특정 기술안을 깊게 다루는 프로젝트형 실습은 "
        "참여 장벽이 생길 수 있고, 반대로 너무 가벼운 wrap-up 활동은 고객사 입장에서 교육 성과가 약하게 보일 수 있습니다."
    )
    add_body(
        doc,
        "또한 고객사는 AI 교육이 앞서 진행되는 반도체 시장 트렌드와 반도체 공정 이해 모듈과 자연스럽게 이어지기를 기대하고 있습니다. "
        "이 점을 고려하면 AI 세션은 새로운 것을 만드는 시간이 아니라, 이미 학습한 내용을 구조화하고 설명 가능하게 만들며 "
        "협업에 필요한 공통 이해를 정리하는 랩업형 실습으로 가는 것이 적절합니다."
    )

    add_heading(doc, "2. 설계 원칙")
    for item in [
        "AI는 기능 시연보다 학습 내용의 구조화와 현업 번역의 도구로 다룬다.",
        "반도체 시장, 기술, 공정 사전 학습과 끊기지 않게 이어간다.",
        "부서와 직무가 달라도 함께 기여할 수 있는 팀 과제로 설계한다.",
        "결과물은 발표용 장식물이 아니라 이후에도 다시 꺼내볼 수 있는 학습 자산으로 남긴다.",
        "마지막에는 협업 메시지를 별도로 한 번 더 강화해 과정 전체를 닫는다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 제안 방향")
    add_body(
        doc,
        "권장 메인 과제는 '반도체 협업 인사이트 팩 제작'입니다. 팀별로 사전 학습한 반도체 시장/공정/기술 주제 중 "
        "하나를 선택해, 타직무자나 협업 파트너도 이해할 수 있는 형태로 재구성하는 과제입니다."
    )
    add_label_paragraph(doc, "핵심 목표", "배운 반도체 내용을 더 잘 이해하고, 더 쉽게 설명하고, 협업에 다시 연결할 수 있게 만든다.")
    add_label_paragraph(doc, "메인 도구", "NotebookLM")
    add_label_paragraph(doc, "성과 방향", "기술 구현보다 공통 이해와 설명력을 강화하는 결과물을 만든다.")

    add_heading(doc, "4. 3일 과정 내 권장 위치")
    schedule = doc.add_table(rows=1, cols=5)
    schedule.style = "Table Grid"
    schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    schedule.autofit = False
    widths = [Cm(1.8), Cm(2.6), Cm(4.1), Cm(4.7), Cm(3.0)]
    headers = ["Day", "모듈", "주요 내용", "AI 연계 포인트", "권장 산출물"]
    for idx, (w, h) in enumerate(zip(widths, headers)):
        schedule.columns[idx].width = w
        cell = schedule.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, color=PRIMARY, size=9.8)

    schedule_rows = [
        ["Day 1", "Module 1", "팀 빌딩", "협업의 기준과 공통 이해의 중요성 심기", "협업 질문 메모"],
        ["Day 1", "Module 2", "반도체 시장 트렌드", "현업에서 의미 있다고 느낀 이슈 기록", "주제 후보 메모"],
        ["Day 2", "Module 3", "반도체 공정 이해", "타직무자가 헷갈릴 개념 포착", "핵심 개념 메모"],
        ["Day 2", "Module 4", "생성형 AI 개념 이해 + 조별 과제 설계", "AI 기본 2시간 이상 + 팀 주제 선정", "팀 과제 설계안"],
        ["Day 3", "Module 5", "생성형 AI Skill 활용", "NotebookLM으로 학습 자산 제작", "인사이트 팩"],
        ["Day 3", "Module 6", "Wrap-up", "AI 발표 후 협업 액티비티와 엔딩 장표로 마무리", "발표, 엔딩 장표"],
    ]
    for row in schedule_rows:
        cells = schedule.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.1)

    add_heading(doc, "5. AI 모듈 상세 제안")
    add_heading(doc, "5-1. Module 4 - AI 기본 이해 및 팀 과제 설계", level=2)
    add_label_paragraph(doc, "권장 시간", "약 2시간 30분")
    add_body(
        doc,
        "Module 4에서는 AI 기본기를 별도로 확보하되, 툴 소개로 흩어지지 않도록 바로 팀 과제 설계까지 연결하는 흐름이 중요합니다. "
        "참여자는 생성형 AI의 기본 원리, 잘하는 일과 한계, 프롬프트의 기본 구조, 검증의 필요성을 이해한 뒤 "
        "NotebookLM 기반 팀 과제를 설계합니다."
    )
    for item in [
        "AI 기본 이해: 생성형 AI의 원리, 강점, 한계, 사실 검증 기준",
        "NotebookLM 소개: 소스 업로드, 질문, Briefing Document, Mind Map, Data Table, Slide Deck, Infographic 등",
        "팀 과제 설계: 주제 선정, 대상 청중 정의, 핵심 질문 3-5개 도출",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5-2. Module 5 - NotebookLM 기반 학습 자산 제작", level=2)
    add_label_paragraph(doc, "권장 시간", "약 2시간 30분")
    add_body(
        doc,
        "Module 5의 핵심은 NotebookLM의 기능을 많이 써보는 것이 아니라, 하나의 반도체 주제를 여러 형태의 학습 자산으로 "
        "재구성하는 것입니다. 팀별로 자료를 업로드한 뒤 보고서, 구조도, 비교표, 발표용 시각자료를 만들어 "
        "학습 내용을 현업 언어로 바꿉니다."
    )

    artifact_table = doc.add_table(rows=1, cols=4)
    artifact_table.style = "Table Grid"
    artifact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    artifact_table.autofit = False
    art_widths = [Cm(3.1), Cm(3.8), Cm(5.2), Cm(4.1)]
    art_headers = ["구분", "권장 도구", "역할", "비고"]
    for idx, (w, h) in enumerate(zip(art_widths, art_headers)):
        artifact_table.columns[idx].width = w
        cell = artifact_table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, color=PRIMARY, size=9.6)

    art_rows = [
        ["설명용", "Briefing Document", "주제 요약, 왜 중요한가, 현업 의미 정리", "필수"],
        ["구조용", "Mind Map 또는 Data Table", "개념 구조와 비교 포인트 시각화", "필수 1개 이상"],
        ["공유용", "Slide Deck 또는 Infographic", "발표 및 팀 공유용 시각 자료", "필수 1개"],
        ["학습보조", "FAQ, Study Guide, Flashcards, Quiz", "복습과 이해도 강화", "선택"],
        ["확장형", "Audio Overview 또는 Video Overview", "빠른 재학습 및 공유", "선택"],
    ]
    for row in art_rows:
        cells = artifact_table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.0)

    add_heading(doc, "5-3. Module 6 - 발표, 협업 액티비티, 엔딩 장표", level=2)
    add_label_paragraph(doc, "권장 시간", "약 1시간 20분 내외")
    add_body(
        doc,
        "Module 6는 AI 발표 세션으로만 끝내기보다, 발표 이후 협업 메시지를 다시 한 번 남기는 흐름으로 설계하는 것이 좋습니다. "
        "권장 순서는 '팀별 AI 결과물 발표 -> 협업 게임 또는 협업 액티비티 -> AI를 활용한 팀별 엔딩 장표 제작 -> 전체 마무리'입니다."
    )
    for item in [
        "팀 발표: 무엇을 이해하게 했는가, 누가 이 결과물을 써야 하는가, 협업에 어떤 도움이 되는가 중심",
        "협업 액티비티: 정보 비대칭 게임, 설명 릴레이 등 15-20분 내외",
        "엔딩 장표: 우리 팀의 반도체 인사이트 한 줄, 협업 메시지 한 줄, 앞으로의 AI 활용 한 줄",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 팀 과제 예시")
    add_body(
        doc,
        "주제는 너무 넓지 않게 잡고, 타직무자가 왜 알아야 하는지가 드러나는 것이 중요합니다. 아래는 실제 운영 가능한 예시입니다."
    )
    example_table = doc.add_table(rows=1, cols=4)
    example_table.style = "Table Grid"
    example_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    example_table.autofit = False
    ex_widths = [Cm(4.0), Cm(4.2), Cm(4.0), Cm(4.0)]
    ex_headers = ["주제 예시", "대상 청중", "권장 아티팩트", "핵심 메시지"]
    for idx, (w, h) in enumerate(zip(ex_widths, ex_headers)):
        example_table.columns[idx].width = w
        cell = example_table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, color=PRIMARY, size=9.4)

    ex_rows = [
        ["HBM이 중요한 이유", "타직무 동료, 신입 구성원", "Briefing, Data Table, Slide Deck", "성능 경쟁은 메모리 단품이 아니라 생태계 협업으로 확장된다."],
        ["첨단 패키징이 주목받는 배경", "비전공자, 협업 파트너", "Briefing, Mind Map, Infographic", "후공정 이해가 협업 속도와 품질 대화의 수준을 바꾼다."],
        ["반도체 시장 변화와 협력사 시사점", "영업/지원/협업 직무", "Briefing, Data Table, FAQ", "시장 변화는 기술 이슈가 아니라 협업 방식의 변화와도 연결된다."],
    ]
    for row in ex_rows:
        cells = example_table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.8)

    add_heading(doc, "7. 기대 효과")
    for item in [
        "반도체 사전 학습이 AI 세션과 자연스럽게 연결되어 교육 흐름이 끊기지 않는다.",
        "직무가 달라도 공통으로 기여할 수 있는 과제라 참여 장벽이 낮다.",
        "결과물이 발표용으로 끝나지 않고 이후에도 다시 활용 가능한 학습 자산으로 남는다.",
        "협업 메시지를 별도 액티비티와 엔딩 장표로 한 번 더 정리해 과정 전체의 마무리가 선명해진다.",
        "AI를 단순 검색 도구가 아니라 이해와 설명, 구조화, 공유의 도구로 경험하게 된다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. 최종 제안")
    closing_box = doc.add_table(rows=1, cols=1)
    closing_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    closing_box.autofit = False
    closing_box.columns[0].width = Cm(16.2)
    cell = closing_box.cell(0, 0)
    set_cell_shading(cell, PRIMARY_LIGHT)
    set_cell_text(
        cell,
        "이번 AI 교육은 따로 노는 툴 교육이 아니라, 이미 학습한 반도체 내용을 NotebookLM 기반 결과물로 재가공해 "
        "현업과 협업에 다시 연결하는 랩업 구조로 가는 것이 가장 적절합니다.",
        bold=True,
        color=PRIMARY,
        size=10.8,
    )

    add_body(
        doc,
        "권장 흐름은 Day 2 Module 4에서 AI 기본기를 확보하고 팀 과제를 설계한 뒤, Day 3 Module 5에서 NotebookLM으로 "
        "학습 자산을 제작하고, Module 6에서 발표와 협업 액티비티, 엔딩 장표로 과정을 마무리하는 방식입니다."
    )

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
